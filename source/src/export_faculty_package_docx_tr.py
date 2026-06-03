"""Export the Turkish faculty presentation package as a DOCX report."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PACKAGE_DIR = ROOT / "docs" / "faculty_presentation_package_tr"
FIGURES_DIR = PACKAGE_DIR / "figures"
TABLES_DIR = PACKAGE_DIR / "tables"
OUTPUT_PATH = PACKAGE_DIR / "hoca_sunum_paketi_tr.docx"


DOC_FILES = [
    ("Danışman İçin Kısa Özet", PACKAGE_DIR / "hoca_icin_kisa_ozet_tr.md"),
    ("Teknik Sunum Raporu", PACKAGE_DIR / "teknik_sunum_raporu_tr.md"),
    ("Gelecek Adımlar ve Araştırma Yol Haritası", PACKAGE_DIR / "gelecek_adimlar_ve_arastirma_yol_haritasi_tr.md"),
    ("Görsel ve Tablo Rehberi", PACKAGE_DIR / "gorsel_ve_tablo_rehberi_tr.md"),
    ("Sunum Konuşma Notları", PACKAGE_DIR / "sunum_konusma_notlari_tr.md"),
]


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)


def _replace_inline_markup(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(_replace_inline_markup(text))


def _flush_buffer(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    _add_paragraph(doc, " ".join(item.strip() for item in buffer if item.strip()))
    buffer.clear()


def _add_markdown(doc: Document, path: Path) -> None:
    buffer: list[str] = []
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            _flush_buffer(doc, buffer)
            continue
        if stripped.startswith("```"):
            continue
        if stripped.startswith("#"):
            _flush_buffer(doc, buffer)
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            doc.add_heading(_replace_inline_markup(stripped[level:].strip()), level=level)
            continue
        if stripped.startswith("- "):
            _flush_buffer(doc, buffer)
            _add_paragraph(doc, stripped[2:], style="List Bullet")
            continue
        if re.match(r"^\d+\.\s", stripped):
            _flush_buffer(doc, buffer)
            _add_paragraph(doc, stripped, style="List Number")
            continue
        buffer.append(stripped)
    _flush_buffer(doc, buffer)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Hocaya Sunulacak Türkçe Teknik Sunum Paketi").bold = True
    title.runs[0].font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Multi-dataset, anomaly-first ve session-safe kestirimci bakım araştırma altyapısı").font.size = Pt(14)

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Oluşturulma tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Çıktı klasörü: {PACKAGE_DIR.relative_to(ROOT)}")

    doc.add_page_break()


def _add_contents(doc: Document) -> None:
    doc.add_heading("İçerik Özeti", level=1)
    items = [
        "1. Danışman İçin Kısa Özet",
        "2. Teknik Sunum Raporu",
        "3. Gelecek Adımlar ve Araştırma Yol Haritası",
        "4. Görsel ve Tablo Rehberi",
        "5. Sunum Konuşma Notları",
        "Ek A. Seçili Şekiller",
        "Ek B. Seçili Tablolar",
    ]
    for item in items:
        _add_paragraph(doc, item, style="List Bullet")
    doc.add_page_break()


def _add_figures(doc: Document) -> None:
    doc.add_heading("Ek A - Seçili Şekiller", level=1)
    for image_path in sorted(FIGURES_DIR.glob("*.png")):
        caption = image_path.stem.replace("_", " ")
        doc.add_heading(caption, level=2)
        path_par = doc.add_paragraph()
        path_par.add_run(f"Dosya yolu: {image_path.relative_to(ROOT)}").italic = True
        doc.add_picture(str(image_path), width=Inches(6.6))
        fig_caption = doc.add_paragraph()
        fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_caption.add_run(caption).italic = True
    doc.add_page_break()


def _write_dataframe_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            text = "" if pd.isna(value) else str(value)
            cells[idx].text = text
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph("")


def _add_tables(doc: Document) -> None:
    doc.add_heading("Ek B - Seçili Tablolar", level=1)
    for csv_path in sorted(TABLES_DIR.glob("*.csv")):
        title = csv_path.stem.replace("_", " ")
        df = pd.read_csv(csv_path)
        _write_dataframe_table(doc, df, title)


def build_docx() -> Path:
    doc = Document()
    _configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Sayfa ")
    _add_page_number(footer)

    _add_cover(doc)
    _add_contents(doc)

    for index, (title, path) in enumerate(DOC_FILES, start=1):
        doc.add_heading(f"{index}. {title}", level=1)
        _add_markdown(doc, path)
        doc.add_page_break()

    _add_figures(doc)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENTATION.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.left_margin = Inches(0.5)
    landscape.right_margin = Inches(0.5)
    _add_tables(doc)

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def main() -> None:
    output = build_docx()
    print(output)


if __name__ == "__main__":
    main()
